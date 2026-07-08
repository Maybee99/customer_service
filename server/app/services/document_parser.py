"""
Document parser for PDF and Word (.docx) files.
Extracts plain text from uploaded documents for RAG ingestion.
"""
import os
from ..utils.logger import logger


class DocumentParser:
    """Parse PDF and DOCX files into plain text."""

    # ── Public API ──────────────────────────────────────────

    @staticmethod
    def parse_file(file_path: str, file_type: str) -> str:
        """Dispatch to the correct parser based on file_type."""
        ext = file_type.lower()
        if ext in ('pdf',):
            return DocumentParser.parse_pdf(file_path)
        elif ext in ('docx', 'doc'):
            return DocumentParser.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from a PDF using pdfplumber.

        Raises:
            FileNotFoundError: if file doesn't exist
            ValueError: if PDF is encrypted or empty
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber is required for PDF parsing. "
                "Install it with: pip install pdfplumber"
            )

        text_parts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                if pdf.pages and hasattr(pdf, 'metadata') and pdf.metadata.get('Encrypted'):
                    # Check if we can still extract (some encrypted PDFs allow it)
                    pass

                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text.strip())

                total_chars = sum(len(p) for p in text_parts)
                logger.info(f"[DocumentParser] PDF parsed: {len(pdf.pages)} pages,"
                            f" {total_chars} chars, file={os.path.basename(file_path)}")

                if total_chars == 0:
                    raise ValueError(
                        "文档无文字内容，可能是扫描件或图片型PDF"
                    )

                # Warn if very few chars per page (possible scanned PDF)
                if len(pdf.pages) > 0 and total_chars / len(pdf.pages) < 50:
                    logger.warning(
                        f"[DocumentParser] Low text density ({total_chars} chars / "
                        f"{len(pdf.pages)} pages). Might be a scanned PDF."
                    )

        except Exception as e:
            err_msg = str(e)
            # Detect common error types
            if 'password' in err_msg.lower() or 'encrypt' in err_msg.lower():
                raise ValueError("文档已加密，请解密后重新上传")
            if 'not a valid' in err_msg.lower() or 'not found' in err_msg.lower():
                raise FileNotFoundError(f"文件无效或不存在: {file_path}")
            raise

        return '\n\n'.join(text_parts)

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from a Word .docx file using python-docx.

        Extracts paragraphs and table cell text.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            )

        try:
            doc = Document(file_path)
        except Exception as e:
            err_msg = str(e).lower()
            if 'not a valid' in err_msg or 'corrupt' in err_msg:
                raise ValueError("文档已损坏，无法打开")
            raise ValueError(f"无法打开文档: {str(e)[:200]}")

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(' | '.join(row_texts))

        full_text = '\n\n'.join(text_parts)
        total_chars = len(full_text)
        logger.info(f"[DocumentParser] DOCX parsed: {len(doc.paragraphs)} paras,"
                    f" {len(doc.tables)} tables, {total_chars} chars,"
                    f" file={os.path.basename(file_path)}")

        if total_chars == 0:
            raise ValueError("文档无文字内容")

        return full_text

    @staticmethod
    def detect_scanned_pdf(file_path: str) -> bool:
        """Check if a PDF appears to be a scanned/image-based document.

        Returns True if the average character count per page is very low (< 20 chars/page),
        which strongly indicates a scanned/image PDF without a text layer.
        """
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                total_chars = 0
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        total_chars += len(text.strip())
                pages = len(pdf.pages) or 1
                avg_chars = total_chars / pages
                return avg_chars < 20
        except Exception:
            return False


# Module-level convenience functions
parse_pdf = DocumentParser.parse_pdf
parse_docx = DocumentParser.parse_docx
parse_file = DocumentParser.parse_file
detect_scanned_pdf = DocumentParser.detect_scanned_pdf
